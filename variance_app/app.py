import streamlit as st
import pandas as pd
import anthropic
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="Variance Commentary", page_icon="📊", layout="wide")

st.markdown("""
<style>
.main { padding-top: 1rem; }
.stApp { background: #fafafa; }
div[data-testid="metric-container"] {
    background: white;
    border: 1px solid #e8e8e8;
    border-radius: 8px;
    padding: 12px 16px;
}
.commentary-box {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 20px;
    margin-top: 8px;
    line-height: 1.8;
    font-size: 14px;
}
.section-tag {
    font-size: 11px;
    font-weight: 600;
    color: #64748b;
    text-transform: uppercase;
    letter-spacing: 0.05em;
    margin-bottom: 6px;
}
.material-row { background: #fffbeb !important; }
.fav-badge { color: #166534; background: #dcfce7; padding: 2px 8px; border-radius: 4px; font-size: 12px; }
.adv-badge { color: #991b1b; background: #fee2e2; padding: 2px 8px; border-radius: 4px; font-size: 12px; }
</style>
""", unsafe_allow_html=True)

def fmt_k(n):
    sign = "+" if n > 0 else ""
    return f"{sign}£{n:,.0f}k" if n != 0 else "—"

def load_sample_tb():
    return pd.DataFrame([
        {"Line item": "Interest income",       "Category": "Revenue",          "Current (£k)": 4820, "Prior (£k)": 3950},
        {"Line item": "Fee income",             "Category": "Revenue",          "Current (£k)": 1240, "Prior (£k)": 1180},
        {"Line item": "FX gains / (losses)",    "Category": "Revenue",          "Current (£k)": -85,  "Prior (£k)": 42},
        {"Line item": "Staff costs",            "Category": "Operating costs",  "Current (£k)": -2100,"Prior (£k)": -1980},
        {"Line item": "Technology & data",      "Category": "Operating costs",  "Current (£k)": -380, "Prior (£k)": -310},
        {"Line item": "Premises & facilities",  "Category": "Operating costs",  "Current (£k)": -145, "Prior (£k)": -145},
        {"Line item": "Professional fees",      "Category": "Operating costs",  "Current (£k)": -290, "Prior (£k)": -210},
        {"Line item": "Regulatory levies",      "Category": "Operating costs",  "Current (£k)": -95,  "Prior (£k)": -88},
        {"Line item": "Intercompany charges",   "Category": "Intercompany",     "Current (£k)": -620, "Prior (£k)": -480},
        {"Line item": "Tax charge",             "Category": "Tax",              "Current (£k)": -580, "Prior (£k)": -490},
    ])

def load_sample_gl():
    return pd.DataFrame([
        {"Date": "2026-05-03", "Line item": "Interest income",     "Reference": "JNL-0501", "Description": "Monthly interest accrual — interco loan to Haarby GmbH (£180m @ 4.8%)",        "Amount (£k)": 720},
        {"Date": "2026-05-03", "Line item": "Interest income",     "Reference": "JNL-0502", "Description": "Monthly interest accrual — external lending portfolio (rate increase +25bp)",   "Amount (£k)": 150},
        {"Date": "2026-05-15", "Line item": "Interest income",     "Reference": "JNL-0515", "Description": "New facility drawdown — Meridian Capital Ltd (£50m @ 5.1%, effective 1 May)",  "Amount (£k)": 213},
        {"Date": "2026-05-31", "Line item": "Interest income",     "Reference": "ADJ-0531", "Description": "Prior period interest accrual catch-up — Q1 underprovision corrected",          "Amount (£k)": -213},
        {"Date": "2026-05-08", "Line item": "Fee income",          "Reference": "JNL-0508", "Description": "Advisory fee — Meridian Capital restructuring mandate (one-off)",               "Amount (£k)": 85},
        {"Date": "2026-05-12", "Line item": "Fee income",          "Reference": "JNL-0512", "Description": "Facility arrangement fee — new revolving credit facility",                     "Amount (£k)": 40},
        {"Date": "2026-05-20", "Line item": "Fee income",          "Reference": "ADJ-0520", "Description": "Accrual release — Hartley Group mandate completed, final invoice raised",      "Amount (£k)": -65},
        {"Date": "2026-05-28", "Line item": "FX gains / (losses)","Reference": "JNL-0528", "Description": "EUR/GBP retranslation loss — EUR-denominated interco balances (spot move)",    "Amount (£k)": -127},
        {"Date": "2026-05-31", "Line item": "FX gains / (losses)","Reference": "ADJ-0531", "Description": "FX hedge MTM gain — cross-currency swap (partial offset)",                     "Amount (£k)": 42},
        {"Date": "2026-05-31", "Line item": "Staff costs",         "Reference": "PAY-0531", "Description": "Monthly payroll — 3 new hires (2x credit analysts, 1x compliance, started 1 May)", "Amount (£k)": -95},
        {"Date": "2026-05-31", "Line item": "Staff costs",         "Reference": "PAY-0532", "Description": "Annual bonus accrual true-up — Q2 performance revision upward",                "Amount (£k)": -25},
        {"Date": "2026-05-15", "Line item": "Technology & data",   "Reference": "INV-0215", "Description": "Bloomberg terminal uplift — 4 additional licences (annualised £48k pa)",      "Amount (£k)": -48},
        {"Date": "2026-05-20", "Line item": "Technology & data",   "Reference": "INV-0220", "Description": "FactSet data subscription renewal — 12% price increase on prior year",        "Amount (£k)": -22},
        {"Date": "2026-05-10", "Line item": "Professional fees",   "Reference": "INV-0310", "Description": "Deloitte audit fee accrual — FY2025 statutory audit (final invoice)",         "Amount (£k)": -55},
        {"Date": "2026-05-22", "Line item": "Professional fees",   "Reference": "INV-0322", "Description": "Linklaters legal fees — Meridian facility documentation",                     "Amount (£k)": -25},
        {"Date": "2026-05-31", "Line item": "Intercompany charges","Reference": "ICO-0531", "Description": "Group treasury recharge — new transfer pricing policy (effective Jan 2026, Q2 catch-up included)", "Amount (£k)": -140},
        {"Date": "2026-05-31", "Line item": "Tax charge",          "Reference": "TAX-0531", "Description": "Current tax accrual — increased profitability vs prior month",                 "Amount (£k)": -90},
    ])

def build_variance_table(tb_df, threshold):
    df = tb_df.copy()
    df["Variance (£k)"] = df["Current (£k)"] - df["Prior (£k)"]
    df["Material"] = df["Variance (£k)"].abs() >= threshold
    df["Direction"] = df["Variance (£k)"].apply(lambda x: "Fav ▲" if x > 0 else ("Adv ▼" if x < 0 else "—"))
    return df

def generate_commentary(tb_df, gl_df, entity, period, basis, threshold, tone, client):
    variance_df = build_variance_table(tb_df, threshold)
    tot_curr = tb_df["Current (£k)"].sum()
    tot_prior = tb_df["Prior (£k)"].sum()
    tot_var = tot_curr - tot_prior

    material_lines = variance_df[variance_df["Material"]]
    immaterial_lines = variance_df[~variance_df["Material"]]

    tb_summary = "\n".join([
        f"  {r['Line item']} ({r['Category']}): current £{r['Current (£k)']}k, prior £{r['Prior (£k)']}k, variance {'+' if r['Variance (£k)']>=0 else ''}{r['Variance (£k)']}k {'[MATERIAL]' if r['Material'] else ''}"
        for _, r in variance_df.iterrows()
    ])

    gl_summary = ""
    if gl_df is not None and not gl_df.empty:
        gl_summary = "\n\nGENERAL LEDGER TRANSACTION DETAIL (use this to explain the WHY behind each movement):\n"
        for line in variance_df[variance_df["Material"]]["Line item"].tolist():
            line_txns = gl_df[gl_df["Line item"] == line] if "Line item" in gl_df.columns else pd.DataFrame()
            if not line_txns.empty:
                gl_summary += f"\n{line}:\n"
                for _, t in line_txns.iterrows():
                    gl_summary += f"  {t.get('Date','')} | {t.get('Reference','')} | {t.get('Description','')} | £{t.get('Amount (£k)','')}k\n"

    tone_instructions = {
        "Audit pack": "Write for inclusion in an audit pack. Precise, factual, passive voice acceptable. Reference line items by exact name. No superlatives. Auditors will verify every number.",
        "Board report": "Write for a board report. Senior audience, clear and concise. Explain commercial significance. Professional but accessible — no jargon.",
        "Regulator": "Write for a regulatory submission. Formal, conservative, complete. Note intercompany items explicitly and explain pricing basis. Avoid ambiguity. Every statement must be defensible."
    }

    prompt = f"""You are a legal entity controller preparing monthly variance commentary for inclusion in a finance reporting pack.

ENTITY: {entity}
PERIOD: {period} ({basis})
MATERIALITY THRESHOLD: £{threshold}k
Net income: current £{tot_curr:.0f}k, prior £{tot_prior:.0f}k, total variance {'+' if tot_var>=0 else ''}£{tot_var:.0f}k

TRIAL BALANCE VARIANCES:
{tb_summary}
{gl_summary}

TONE: {tone_instructions.get(tone, tone_instructions['Audit pack'])}

INSTRUCTIONS:
- For material lines, explain the SPECIFIC DRIVER using the GL transaction detail where available — not just the direction and amount
- Reference specific journal references, counterparties, or commercial events from the GL where relevant
- For immaterial lines, group them concisely in a single sentence
- Use £k throughout
- Do not start with "I" or "The entity"
- Write in continuous prose paragraphs, not bullet points

Return a JSON object with exactly these keys:
{{
  "overview": "2-3 sentence paragraph: total net income movement and the primary 1-2 drivers",
  "material": "paragraph per material line or grouped where related — explain the WHY using GL detail",
  "immaterial": "single sentence covering all immaterial lines and their aggregate movement"
}}

Return ONLY the JSON. No markdown fences. No preamble."""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    import json, re
    raw = response.content[0].text
    match = re.search(r'\{[\s\S]*\}', raw)
    if not match:
        raise ValueError("Could not parse JSON from response")
    return json.loads(match.group(0))

def export_to_word(entity, period, basis, commentary, variance_df):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    title = doc.add_heading(f"Variance Commentary — {entity}", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1c, 0x2b, 0x3a)

    sub = doc.add_paragraph(f"{period} · {basis} · Prepared {datetime.now().strftime('%d %b %Y')}")
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    sub.runs[0].font.size = Pt(10)

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(commentary.get("overview", ""))

    doc.add_heading("2. Material variances", level=2)
    doc.add_paragraph(commentary.get("material", ""))

    doc.add_heading("3. Immaterial items", level=2)
    doc.add_paragraph(commentary.get("immaterial", ""))

    doc.add_heading("4. Variance summary", level=2)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Line item", "Category", "Current (£k)", "Prior (£k)", "Variance (£k)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for _, row in variance_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = row["Line item"]
        cells[1].text = row["Category"]
        cells[2].text = f"{row['Current (£k)']:,.0f}"
        cells[3].text = f"{row['Prior (£k)']:,.0f}"
        v = row["Variance (£k)"]
        cells[4].text = f"{'+' if v>=0 else ''}{v:,.0f}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### Configuration")
    api_key = st.text_input("Anthropic API key", type="password", help="Your key is never stored")
    entity = st.text_input("Legal entity", value="Haarby Capital Ltd (UK)")
    period = st.text_input("Period", value="May 2026")
    basis = st.selectbox("Comparison basis", ["vs prior month", "vs budget", "vs prior year"])
    threshold = st.number_input("Materiality threshold (£k)", value=50, step=10)
    tone = st.selectbox("Commentary tone", ["Audit pack", "Board report", "Regulator"])
    st.markdown("---")
    st.markdown("**How to use**")
    st.markdown("1. Enter your API key\n2. Upload TB + GL, or load sample data\n3. Generate commentary\n4. Export to Word")

# ── Main ──────────────────────────────────────────────────────────────────────
st.markdown("## Variance Commentary")
st.markdown(f"**{entity}** · {period} · {basis}")

col1, col2 = st.columns(2)

with col1:
    st.markdown("#### Trial balance")
    tb_file = st.file_uploader("Upload TB (Excel or CSV)", type=["xlsx","xls","csv"], key="tb")
    if st.button("Load sample TB"):
        st.session_state["tb_df"] = load_sample_tb()
    if tb_file:
        try:
            st.session_state["tb_df"] = pd.read_csv(tb_file) if tb_file.name.endswith(".csv") else pd.read_excel(tb_file)
        except Exception as e:
            st.error(f"Could not read TB file: {e}")

with col2:
    st.markdown("#### General ledger")
    gl_file = st.file_uploader("Upload GL transactions (Excel or CSV)", type=["xlsx","xls","csv"], key="gl")
    if st.button("Load sample GL"):
        st.session_state["gl_df"] = load_sample_gl()
    if gl_file:
        try:
            st.session_state["gl_df"] = pd.read_csv(gl_file) if gl_file.name.endswith(".csv") else pd.read_excel(gl_file)
        except Exception as e:
            st.error(f"Could not read GL file: {e}")

tb_df = st.session_state.get("tb_df")
gl_df = st.session_state.get("gl_df")

if tb_df is not None:
    st.markdown("---")
    variance_df = build_variance_table(tb_df, threshold)
    tot_curr = tb_df["Current (£k)"].sum()
    tot_prior = tb_df["Prior (£k)"].sum()
    tot_var = tot_curr - tot_prior
    mat_count = variance_df["Material"].sum()

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Net income (current)", f"£{tot_curr:,.0f}k")
    m2.metric("Net income (prior)", f"£{tot_prior:,.0f}k")
    m3.metric("Total variance", f"{'+' if tot_var>=0 else ''}£{tot_var:,.0f}k")
    m4.metric("Material lines", int(mat_count))

    st.markdown("#### Variance table")

    display_df = variance_df[["Line item","Category","Current (£k)","Prior (£k)","Variance (£k)","Direction","Material"]].copy()

    def highlight_material(row):
        if row["Material"]:
            return ["background-color: #fffbeb"] * len(row)
        return [""] * len(row)

    styled = (display_df.drop(columns=["Material"])
              .style.apply(highlight_material, axis=1, subset=pd.IndexSlice[:, display_df.drop(columns=["Material"]).columns])
              .format({"Current (£k)": "{:,.0f}", "Prior (£k)": "{:,.0f}", "Variance (£k)": "{:+,.0f}"}))
    st.dataframe(styled, use_container_width=True, hide_index=True)

    if gl_df is not None:
        with st.expander(f"GL transactions loaded — {len(gl_df)} rows"):
            st.dataframe(gl_df, use_container_width=True, hide_index=True)

    st.markdown("---")

    col_gen, col_note = st.columns([1, 3])
    with col_gen:
        generate = st.button("Generate commentary", type="primary", use_container_width=True)
    with col_note:
        if gl_df is None:
            st.info("No GL data loaded — commentary will explain amounts only. Upload GL transactions for driver-level insight.")
        else:
            st.success(f"GL data loaded — commentary will explain specific drivers from {len(gl_df)} transactions.")

    if generate:
        if not api_key:
            st.error("Enter your Anthropic API key in the sidebar.")
        else:
            with st.spinner("Analysing variances and GL transactions…"):
                try:
                    client = anthropic.Anthropic(api_key=api_key)
                    result = generate_commentary(tb_df, gl_df, entity, period, basis, threshold, tone, client)
                    st.session_state["commentary"] = result
                    st.session_state["variance_df_export"] = variance_df
                except Exception as e:
                    st.error(f"Error: {e}")

    commentary = st.session_state.get("commentary")
    if commentary:
        st.markdown("### Generated commentary")
        st.markdown(f"<div class='section-tag'>Overview</div><div class='commentary-box'>{commentary.get('overview','')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='section-tag' style='margin-top:16px'>Material variances</div><div class='commentary-box'>{commentary.get('material','')}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='section-tag' style='margin-top:16px'>Immaterial items</div><div class='commentary-box'>{commentary.get('immaterial','')}</div>", unsafe_allow_html=True)

        st.markdown("---")
        word_buf = export_to_word(entity, period, basis, commentary, st.session_state["variance_df_export"])
        st.download_button(
            label="Download Word doc",
            data=word_buf,
            file_name=f"variance_commentary_{period.replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Load sample data or upload your trial balance to get started.")
